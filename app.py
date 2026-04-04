from flask import Flask, request, send_file, jsonify
from flask_cors import CORS
from pypdf import PdfWriter, PdfReader
from docx import Document as DocxDocument
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from PIL import Image
import os, tempfile, uuid, io, subprocess, shutil, platform, requests as req_lib

app = Flask(__name__)
CORS(app)

TMP = tempfile.gettempdir()

def tmp_path(ext):
    return os.path.join(TMP, f"{uuid.uuid4()}.{ext}")

def cleanup(*paths):
    for p in paths:
        try:
            if p and os.path.exists(p): os.remove(p)
        except: pass

def get_soffice_path():
    if platform.system() == 'Windows':
        for p in [r'C:\Program Files\LibreOffice\program\soffice.exe',
                  r'C:\Program Files (x86)\LibreOffice\program\soffice.exe']:
            if os.path.exists(p): return p
        raise RuntimeError("LibreOffice not found")
    else:
        for cmd in ['libreoffice','soffice']:
            if shutil.which(cmd): return cmd
        raise RuntimeError("LibreOffice not installed on server")

def libreoffice_convert(input_path, output_format, output_dir):
    soffice = get_soffice_path()
    env = os.environ.copy()
    if platform.system() != 'Windows':
        env['HOME'] = output_dir
    cmd = [soffice,'--headless','--norestore','--nofirststartwizard',
           '--convert-to', output_format,'--outdir', output_dir, input_path]
    result = subprocess.run(cmd, capture_output=True, text=True, timeout=60, env=env)
    if result.returncode != 0:
        raise RuntimeError(f"LibreOffice error: {result.stderr or result.stdout}")
    base = os.path.splitext(os.path.basename(input_path))[0]
    ext = output_format.split(':')[0]
    out_file = os.path.join(output_dir, f"{base}.{ext}")
    if not os.path.exists(out_file):
        files = os.listdir(output_dir)
        matches = [f for f in files if f.startswith(base)]
        if matches: out_file = os.path.join(output_dir, matches[0])
        else: raise RuntimeError(f"Output file not found. Files: {files}")
    return out_file

# ── GEMINI SETUP ──────────────────────────────────────────────
GEMINI_KEY = os.environ.get('GEMINI_API_KEY', '')
# ✅ FIXED: Using gemini-2.0-flash (1.5-flash is deprecated/404)
GEMINI_URL = 'https://generativelanguage.googleapis.com/v1beta/models/gemini-2.0-flash:generateContent'

def call_gemini(prompt):
    headers = {'Content-Type': 'application/json'}
    data = {'contents': [{'parts': [{'text': prompt}]}]}
    r = req_lib.post(f'{GEMINI_URL}?key={GEMINI_KEY}', headers=headers, json=data, timeout=40)
    r.raise_for_status()
    return r.json()['candidates'][0]['content']['parts'][0]['text']

# ── HEALTH ────────────────────────────────────────────────────
@app.route('/')
def health():
    try:
        soffice = get_soffice_path()
        lo_status = f"Found: {soffice}"
    except Exception as e:
        lo_status = f"NOT found: {str(e)}"
    return jsonify({
        'status': 'ok',
        'message': 'FreeToolKit API running',
        'libreoffice': lo_status,
        'gemini_key_set': bool(GEMINI_KEY),
        'tools': ['word-to-pdf','pdf-to-word','pdf-merge','pdf-split','img-to-pdf','ai-chat','ai-resume-build','ai-resume-improve']
    })

# ── 1. WORD TO PDF ────────────────────────────────────────────
@app.route('/word-to-pdf', methods=['POST'])
def word_to_pdf():
    if 'file' not in request.files:
        return jsonify({'error': 'File not found'}), 400
    file = request.files['file']
    if not file.filename.lower().endswith(('.doc','.docx')):
        return jsonify({'error': 'DOC or DOCX only'}), 400
    input_path = tmp_path('docx')
    out_dir = tempfile.mkdtemp()
    file.save(input_path)
    try:
        out_file = libreoffice_convert(input_path, 'pdf', out_dir)
        return send_file(out_file, mimetype='application/pdf', as_attachment=True, download_name='converted.pdf')
    except Exception as e:
        return jsonify({'error': str(e)}), 500
    finally:
        cleanup(input_path)
        shutil.rmtree(out_dir, ignore_errors=True)

# ── 2. PDF TO WORD ────────────────────────────────────────────
@app.route('/pdf-to-word', methods=['POST'])
def pdf_to_word():
    if 'file' not in request.files:
        return jsonify({'error': 'File not found'}), 400
    file = request.files['file']
    input_path = tmp_path('pdf')
    file.save(input_path)
    try:
        reader = PdfReader(input_path)
        doc = DocxDocument()
        doc.add_heading('Converted Document', 0)
        for i, page in enumerate(reader.pages):
            text = page.extract_text()
            if text and text.strip():
                doc.add_heading(f'Page {i+1}', level=2)
                for line in text.split('\n'):
                    if line.strip(): doc.add_paragraph(line.strip())
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return send_file(buf, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                        as_attachment=True, download_name='converted.docx')
    except Exception as e:
        return jsonify({'error': str(e)}), 500
    finally:
        cleanup(input_path)

# ── 3. PDF MERGE ──────────────────────────────────────────────
@app.route('/pdf-merge', methods=['POST'])
def pdf_merge():
    files = request.files.getlist('files')
    if len(files) < 2:
        return jsonify({'error': 'At least 2 PDFs required'}), 400
    saved = []
    try:
        writer = PdfWriter()
        for file in files:
            path = tmp_path('pdf')
            file.save(path)
            saved.append(path)
            for page in PdfReader(path).pages:
                writer.add_page(page)
        buf = io.BytesIO()
        writer.write(buf)
        buf.seek(0)
        return send_file(buf, mimetype='application/pdf', as_attachment=True, download_name='merged.pdf')
    except Exception as e:
        return jsonify({'error': str(e)}), 500
    finally:
        cleanup(*saved)

# ── 4. PDF SPLIT ──────────────────────────────────────────────
@app.route('/pdf-split', methods=['POST'])
def pdf_split():
    if 'file' not in request.files:
        return jsonify({'error': 'File not found'}), 400
    file = request.files['file']
    pages_param = request.form.get('pages','').strip()
    input_path = tmp_path('pdf')
    file.save(input_path)
    try:
        reader = PdfReader(input_path)
        total = len(reader.pages)
        pages_to_extract = parse_pages(pages_param, total)
        writer = PdfWriter()
        for i in pages_to_extract:
            writer.add_page(reader.pages[i])
        buf = io.BytesIO()
        writer.write(buf)
        buf.seek(0)
        return send_file(buf, mimetype='application/pdf', as_attachment=True, download_name='split.pdf')
    except Exception as e:
        return jsonify({'error': str(e)}), 500
    finally:
        cleanup(input_path)

# ── 5. IMAGE TO PDF ───────────────────────────────────────────
@app.route('/img-to-pdf', methods=['POST'])
def img_to_pdf():
    files = request.files.getlist('files')
    if not files:
        return jsonify({'error': 'No images found'}), 400
    saved = []
    try:
        buf = io.BytesIO()
        c = canvas.Canvas(buf, pagesize=A4)
        w, h = A4
        for file in files:
            path = tmp_path('img')
            file.save(path)
            saved.append(path)
            img = Image.open(path)
            if img.mode in ('RGBA','P','LA'):
                background = Image.new('RGB', img.size, (255,255,255))
                if img.mode == 'P': img = img.convert('RGBA')
                background.paste(img, mask=img.split()[-1] if img.mode=='RGBA' else None)
                img = background
                img.save(path, 'JPEG', quality=95)
            elif img.mode != 'RGB':
                img = img.convert('RGB')
                img.save(path, 'JPEG', quality=95)
            img_w, img_h = img.size
            ratio = min(w/img_w, h/img_h)
            new_w, new_h = img_w*ratio, img_h*ratio
            c.drawImage(path, (w-new_w)/2, (h-new_h)/2, new_w, new_h, preserveAspectRatio=True)
            c.showPage()
        c.save()
        buf.seek(0)
        return send_file(buf, mimetype='application/pdf', as_attachment=True, download_name='images.pdf')
    except Exception as e:
        return jsonify({'error': str(e)}), 500
    finally:
        cleanup(*saved)

# ── 6. PDF COMPRESS ───────────────────────────────────────────
@app.route('/pdf-compress', methods=['POST'])
def pdf_compress():
    if 'file' not in request.files:
        return jsonify({'error': 'File not found'}), 400
    file = request.files['file']
    input_path = tmp_path('pdf')
    file.save(input_path)
    original_size = os.path.getsize(input_path)
    try:
        reader = PdfReader(input_path)
        writer = PdfWriter()
        for page in reader.pages:
            page.compress_content_streams()
            writer.add_page(page)
        if reader.metadata:
            writer.add_metadata(reader.metadata)
        buf = io.BytesIO()
        writer.write(buf)
        buf.seek(0)
        compressed_size = buf.getbuffer().nbytes
        saved_percent = round((1 - compressed_size/original_size)*100, 1)
        response = send_file(buf, mimetype='application/pdf', as_attachment=True, download_name='compressed.pdf')
        response.headers['X-Original-Size'] = str(original_size)
        response.headers['X-Compressed-Size'] = str(compressed_size)
        response.headers['X-Saved-Percent'] = str(saved_percent)
        response.headers['Access-Control-Expose-Headers'] = 'X-Original-Size, X-Compressed-Size, X-Saved-Percent'
        return response
    except Exception as e:
        return jsonify({'error': str(e)}), 500
    finally:
        cleanup(input_path)

# ── 7. AI CHAT ────────────────────────────────────────────────
@app.route('/ai-chat', methods=['POST'])
def ai_chat():
    data = request.get_json()
    if not data:
        return jsonify({'error': 'No data provided'}), 400

    system   = data.get('system', '')
    messages = data.get('messages', [])

    if not messages:
        return jsonify({'error': 'No messages provided'}), 400

    # Build Gemini contents
    contents = []
    if system:
        contents.append({'role':'user','parts':[{'text':f'[SYSTEM]\n{system}\n[/SYSTEM]\n\nAcknowledge.'}]})
        contents.append({'role':'model','parts':[{'text':'Understood! I am your AI Resume Assistant.'}]})

    for msg in messages:
        role = 'user' if msg.get('role') == 'user' else 'model'
        contents.append({'role': role, 'parts': [{'text': msg.get('content','')}]})

    try:
        headers = {'Content-Type': 'application/json'}
        r = req_lib.post(f'{GEMINI_URL}?key={GEMINI_KEY}',
                        headers=headers, json={'contents': contents}, timeout=40)
        if not r.ok:
            return jsonify({'error': f'Gemini API error: {r.status_code} — {r.text[:300]}'}), 500
        reply = r.json()['candidates'][0]['content']['parts'][0]['text']
        return jsonify({'reply': reply})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

# ── 8. AI RESUME BUILD ────────────────────────────────────────
def build_resume_pdf(text):
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4,
        rightMargin=0.65*inch, leftMargin=0.65*inch,
        topMargin=0.65*inch, bottomMargin=0.65*inch)

    name_s = ParagraphStyle('n', fontSize=18, fontName='Helvetica-Bold', alignment=TA_CENTER, spaceAfter=3, textColor=colors.HexColor('#1a1a2e'))
    cont_s = ParagraphStyle('c', fontSize=9, fontName='Helvetica', alignment=TA_CENTER, spaceAfter=10, textColor=colors.HexColor('#555'))
    head_s = ParagraphStyle('h', fontSize=11, fontName='Helvetica-Bold', spaceBefore=8, spaceAfter=3, textColor=colors.HexColor('#5b4fff'))
    body_s = ParagraphStyle('b', fontSize=9.5, fontName='Helvetica', spaceAfter=3, leading=14, textColor=colors.HexColor('#222'))
    bull_s = ParagraphStyle('bl', fontSize=9.5, fontName='Helvetica', spaceAfter=2, leading=13, leftIndent=14, textColor=colors.HexColor('#333'))

    story = []
    lines = [l.rstrip() for l in text.strip().split('\n')]
    name_done = False

    for line in lines:
        if not line.strip():
            story.append(Spacer(1, 3)); continue
        if line.startswith('##'):
            txt = line.replace('##','').replace('**','').strip()
            story.append(HRFlowable(width='100%', thickness=0.5, color=colors.HexColor('#5b4fff'), spaceAfter=4))
            story.append(Paragraph(txt.upper(), head_s))
        elif line.startswith('#') and not name_done:
            story.append(Paragraph(line.replace('#','').replace('**','').strip(), name_s))
            name_done = True
        elif line.startswith('- ') or line.startswith('• '):
            story.append(Paragraph('• ' + line[2:].replace('**','').strip(), bull_s))
        elif any(x in line for x in ['@','|','+91','linkedin','github','http']) and not name_done:
            story.append(Paragraph(line.replace('**','').strip(), cont_s))
            name_done = True
        elif line.startswith('**') and line.endswith('**'):
            story.append(Paragraph(line.replace('**','').strip(),
                ParagraphStyle('sb', fontSize=9.5, fontName='Helvetica-Bold', spaceAfter=2, textColor=colors.HexColor('#222'))))
        else:
            clean = line.replace('**','').replace('##','').replace('#','').strip()
            if clean:
                if not name_done and len(clean) < 50:
                    story.append(Paragraph(clean, name_s)); name_done = True
                else:
                    story.append(Paragraph(clean, body_s))

    doc.build(story)
    buf.seek(0)
    return buf

@app.route('/ai-resume-build', methods=['POST'])
def ai_resume_build():
    data = request.get_json()
    if not data: return jsonify({'error': 'No data'}), 400
    prompt = f"""Create a professional resume in clean plain text for:
Name: {data.get('name','')}
Email: {data.get('email','')}
Phone: {data.get('phone','')}
LinkedIn: {data.get('linkedin','')}
Summary: {data.get('summary','')}
Skills: {data.get('skills','')}
Education: {data.get('education','')}
Experience: {data.get('experience','')}
Projects: {data.get('projects','')}
Certifications: {data.get('certifications','')}

Format: First line = name only. Second line = contact with | separators.
Use ## for sections. Use - for bullets. Use **bold** for job titles.
ATS-friendly. Only output resume text."""
    try:
        resume_text = call_gemini(prompt)
        return send_file(build_resume_pdf(resume_text), mimetype='application/pdf',
                        as_attachment=True, download_name='resume.pdf')
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/ai-resume-improve', methods=['POST'])
def ai_resume_improve():
    if 'file' not in request.files: return jsonify({'error': 'File not found'}), 400
    file = request.files['file']
    job_role = request.form.get('job_role', 'Software Engineer')
    ext = 'pdf' if file.filename.lower().endswith('.pdf') else 'docx'
    input_path = tmp_path(ext)
    file.save(input_path)
    try:
        if ext == 'pdf':
            reader = PdfReader(input_path)
            resume_text = ' '.join([p.extract_text() or '' for p in reader.pages])
        else:
            doc = DocxDocument(input_path)
            resume_text = ' '.join([p.text for p in doc.paragraphs if p.text.strip()])
        prompt = f"""Improve this resume for "{job_role}". Keep all info, improve wording.
Format: First line=name, second line=contact with |, ## for sections, - for bullets, **bold** for titles.
Original: {resume_text}
Output only the improved resume."""
        improved = call_gemini(prompt)
        return send_file(build_resume_pdf(improved), mimetype='application/pdf',
                        as_attachment=True, download_name='improved_resume.pdf')
    except Exception as e:
        return jsonify({'error': str(e)}), 500
    finally:
        cleanup(input_path)

def parse_pages(pages_param, total):
    if not pages_param: return list(range(total))
    pages = set()
    for part in pages_param.split(','):
        part = part.strip()
        if '-' in part:
            s, e = part.split('-')
            for p in range(int(s)-1, int(e)):
                if 0 <= p < total: pages.add(p)
        else:
            p = int(part)-1
            if 0 <= p < total: pages.add(p)
    return sorted(pages)

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    app.run(host='0.0.0.0', port=port)