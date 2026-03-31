from flask import Flask, request, send_file, jsonify
from flask_cors import CORS
from pypdf import PdfWriter, PdfReader
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from PIL import Image
import os, tempfile, uuid, io

app = Flask(__name__)
CORS(app)

TMP = tempfile.gettempdir()

def tmp_path(ext):
    return os.path.join(TMP, f"{uuid.uuid4()}.{ext}")

def cleanup(*paths):
    for p in paths:
        try:
            if p and os.path.exists(p): os.remove(p)
        except: pass

@app.route('/')
def health():
    return jsonify({'status': 'ok', 'message': 'ಕನ್ನಡ PDF API ಚಾಲನೆಯಲ್ಲಿದೆ'})

@app.route('/word-to-pdf', methods=['POST'])
def word_to_pdf():
    if 'file' not in request.files:
        return jsonify({'error': 'ಫೈಲ್ ಕಂಡುಬಂದಿಲ್ಲ'}), 400
    file = request.files['file']
    if not file.filename.lower().endswith(('.doc', '.docx')):
        return jsonify({'error': 'DOC ಅಥವಾ DOCX ಮಾತ್ರ ಸ್ವೀಕಾರ'}), 400
    input_path = tmp_path('docx')
    file.save(input_path)
    try:
        doc = Document(input_path)
        buf = io.BytesIO()
        c = canvas.Canvas(buf, pagesize=A4)
        width, height = A4
        y = height - 50
        c.setFont("Helvetica", 11)
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                y -= 10
                continue
            words = text.split()
            line = ""
            for word in words:
                test = line + " " + word if line else word
                if c.stringWidth(test, "Helvetica", 11) < width - 80:
                    line = test
                else:
                    c.drawString(40, y, line)
                    y -= 18
                    line = word
                    if y < 60:
                        c.showPage()
                        c.setFont("Helvetica", 11)
                        y = height - 50
            if line:
                c.drawString(40, y, line)
                y -= 18
            if y < 60:
                c.showPage()
                c.setFont("Helvetica", 11)
                y = height - 50
        c.save()
        buf.seek(0)
        return send_file(buf, mimetype='application/pdf', as_attachment=True, download_name='converted.pdf')
    except Exception as e:
        return jsonify({'error': str(e)}), 500
    finally:
        cleanup(input_path)

@app.route('/pdf-merge', methods=['POST'])
def pdf_merge():
    files = request.files.getlist('files')
    if len(files) < 2:
        return jsonify({'error': 'ಕನಿಷ್ಠ 2 PDF ಬೇಕು'}), 400
    saved = []
    try:
        writer = PdfWriter()
        for file in files:
            path = tmp_path('pdf')
            file.save(path)
            saved.append(path)
            for page in PdfReader(path).pages:
                writer.add_page(page)
        buf = io.BytesIO()
        writer.write(buf)
        buf.seek(0)
        return send_file(buf, mimetype='application/pdf', as_attachment=True, download_name='merged.pdf')
    except Exception as e:
        return jsonify({'error': str(e)}), 500
    finally:
        cleanup(*saved)

@app.route('/pdf-split', methods=['POST'])
def pdf_split():
    if 'file' not in request.files:
        return jsonify({'error': 'ಫೈಲ್ ಕಂಡುಬಂದಿಲ್ಲ'}), 400
    file = request.files['file']
    pages_param = request.form.get('pages', '').strip()
    input_path = tmp_path('pdf')
    file.save(input_path)
    try:
        reader = PdfReader(input_path)
        total = len(reader.pages)
        pages_to_extract = parse_pages(pages_param, total)
        writer = PdfWriter()
        for i in pages_to_extract:
            writer.add_page(reader.pages[i])
        buf = io.BytesIO()
        writer.write(buf)
        buf.seek(0)
        return send_file(buf, mimetype='application/pdf', as_attachment=True, download_name='split.pdf')
    except Exception as e:
        return jsonify({'error': str(e)}), 500
    finally:
        cleanup(input_path)

@app.route('/img-to-pdf', methods=['POST'])
def img_to_pdf():
    files = request.files.getlist('files')
    if not files:
        return jsonify({'error': 'ಚಿತ್ರಗಳು ಕಂಡುಬಂದಿಲ್ಲ'}), 400
    saved = []
    try:
        buf = io.BytesIO()
        c = canvas.Canvas(buf, pagesize=A4)
        w, h = A4
        for file in files:
            path = tmp_path('img')
            file.save(path)
            saved.append(path)
            img = Image.open(path)
            img_w, img_h = img.size
            ratio = min(w / img_w, h / img_h)
            new_w, new_h = img_w * ratio, img_h * ratio
            c.drawImage(path, (w-new_w)/2, (h-new_h)/2, new_w, new_h)
            c.showPage()
        c.save()
        buf.seek(0)
        return send_file(buf, mimetype='application/pdf', as_attachment=True, download_name='images.pdf')
    except Exception as e:
        return jsonify({'error': str(e)}), 500
    finally:
        cleanup(*saved)

@app.route('/pdf-to-word', methods=['POST'])
def pdf_to_word():
    if 'file' not in request.files:
        return jsonify({'error': 'ಫೈಲ್ ಕಂಡುಬಂದಿಲ್ಲ'}), 400
    file = request.files['file']
    input_path = tmp_path('pdf')
    file.save(input_path)
    try:
        reader = PdfReader(input_path)
        doc = Document()
        doc.add_heading('Converted from PDF', 0)
        for i, page in enumerate(reader.pages):
            text = page.extract_text()
            if text:
                doc.add_heading(f'Page {i+1}', level=2)
                doc.add_paragraph(text)
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return send_file(buf, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document', as_attachment=True, download_name='converted.docx')
    except Exception as e:
        return jsonify({'error': str(e)}), 500
    finally:
        cleanup(input_path)

def parse_pages(pages_param, total):
    if not pages_param:
        return list(range(total))
    pages = set()
    for part in pages_param.split(','):
        part = part.strip()
        if '-' in part:
            s, e = part.split('-')
            for p in range(int(s)-1, int(e)):
                if 0 <= p < total: pages.add(p)
        else:
            p = int(part) - 1
            if 0 <= p < total: pages.add(p)
    return sorted(pages)

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    app.run(host='0.0.0.0', port=port)