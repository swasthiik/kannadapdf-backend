from flask import Flask, request, send_file, jsonify
from flask_cors import CORS
from pypdf import PdfWriter, PdfReader
from docx import Document as DocxDocument
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from PIL import Image
import os, tempfile, uuid, io, subprocess, shutil, platform

app = Flask(__name__)
CORS(app)

TMP = tempfile.gettempdir()

def tmp_path(ext):
    return os.path.join(TMP, f"{uuid.uuid4()}.{ext}")

def cleanup(*paths):
    for p in paths:
        try:
            if p and os.path.exists(p): os.remove(p)
        except: pass

def get_soffice_path():
    if platform.system() == 'Windows':
        paths = [
            r'C:\Program Files\LibreOffice\program\soffice.exe',
            r'C:\Program Files (x86)\LibreOffice\program\soffice.exe',
        ]
        for p in paths:
            if os.path.exists(p):
                return p
        raise RuntimeError("LibreOffice not found")
    else:
        for cmd in ['libreoffice', 'soffice']:
            if shutil.which(cmd):
                return cmd
        raise RuntimeError("LibreOffice not installed on server")

def libreoffice_convert(input_path, output_format, output_dir):
    soffice = get_soffice_path()
    env = os.environ.copy()
    if platform.system() != 'Windows':
        env['HOME'] = output_dir

    cmd = [
        soffice, '--headless', '--norestore', '--nofirststartwizard',
        '--convert-to', output_format,
        '--outdir', output_dir,
        input_path
    ]
    result = subprocess.run(cmd, capture_output=True, text=True, timeout=60, env=env)

    if result.returncode != 0:
        raise RuntimeError(f"LibreOffice error: {result.stderr or result.stdout}")

    base = os.path.splitext(os.path.basename(input_path))[0]
    ext = output_format.split(':')[0]
    out_file = os.path.join(output_dir, f"{base}.{ext}")

    if not os.path.exists(out_file):
        files = os.listdir(output_dir)
        matches = [f for f in files if f.startswith(base)]
        if matches:
            out_file = os.path.join(output_dir, matches[0])
        else:
            raise RuntimeError(f"Output file not found. Files: {files}")
    return out_file


# ─────────────────────────────────────────
# HEALTH CHECK
# ─────────────────────────────────────────
@app.route('/')
def health():
    try:
        soffice = get_soffice_path()
        lo_status = f"Found: {soffice}"
    except Exception as e:
        lo_status = f"NOT found: {str(e)}"
    return jsonify({
        'status': 'ok',
        'message': 'ಕನ್ನಡ PDF API ಚಾಲನೆಯಲ್ಲಿದೆ',
        'libreoffice': lo_status,
        'tools': ['word-to-pdf', 'pdf-to-word', 'pdf-merge', 'pdf-split', 'img-to-pdf', 'pdf-compress']
    })


# ─────────────────────────────────────────
# 1. WORD TO PDF — LibreOffice (perfect formatting)
# ─────────────────────────────────────────
@app.route('/word-to-pdf', methods=['POST'])
def word_to_pdf():
    if 'file' not in request.files:
        return jsonify({'error': 'ಫೈಲ್ ಕಂಡುಬಂದಿಲ್ಲ'}), 400
    file = request.files['file']
    if not file.filename.lower().endswith(('.doc', '.docx')):
        return jsonify({'error': 'DOC ಅಥವಾ DOCX ಮಾತ್ರ ಸ್ವೀಕಾರ'}), 400

    input_path = tmp_path('docx')
    out_dir = tempfile.mkdtemp()
    file.save(input_path)
    try:
        out_file = libreoffice_convert(input_path, 'pdf', out_dir)
        return send_file(out_file, mimetype='application/pdf', as_attachment=True, download_name='converted.pdf')
    except Exception as e:
        return jsonify({'error': str(e)}), 500
    finally:
        cleanup(input_path)
        shutil.rmtree(out_dir, ignore_errors=True)


# ─────────────────────────────────────────
# 2. PDF TO WORD — extract text into docx
# ─────────────────────────────────────────
@app.route('/pdf-to-word', methods=['POST'])
def pdf_to_word():
    if 'file' not in request.files:
        return jsonify({'error': 'ಫೈಲ್ ಕಂಡುಬಂದಿಲ್ಲ'}), 400
    file = request.files['file']
    if not file.filename.lower().endswith('.pdf'):
        return jsonify({'error': 'PDF ಮಾತ್ರ ಸ್ವೀಕಾರ'}), 400

    input_path = tmp_path('pdf')
    file.save(input_path)
    try:
        reader = PdfReader(input_path)
        doc = DocxDocument()
        doc.add_heading('Converted Document', 0)

        for i, page in enumerate(reader.pages):
            text = page.extract_text()
            if text and text.strip():
                doc.add_heading(f'Page {i+1}', level=2)
                for line in text.split('\n'):
                    if line.strip():
                        doc.add_paragraph(line.strip())

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return send_file(
            buf,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name='converted.docx'
        )
    except Exception as e:
        return jsonify({'error': str(e)}), 500
    finally:
        cleanup(input_path)


# ─────────────────────────────────────────
# 3. PDF MERGE
# ─────────────────────────────────────────
@app.route('/pdf-merge', methods=['POST'])
def pdf_merge():
    files = request.files.getlist('files')
    if len(files) < 2:
        return jsonify({'error': 'ಕನಿಷ್ಠ 2 PDF ಬೇಕು'}), 400

    saved = []
    try:
        writer = PdfWriter()
        for file in files:
            path = tmp_path('pdf')
            file.save(path)
            saved.append(path)
            for page in PdfReader(path).pages:
                writer.add_page(page)
        buf = io.BytesIO()
        writer.write(buf)
        buf.seek(0)
        return send_file(buf, mimetype='application/pdf', as_attachment=True, download_name='merged.pdf')
    except Exception as e:
        return jsonify({'error': str(e)}), 500
    finally:
        cleanup(*saved)


# ─────────────────────────────────────────
# 4. PDF SPLIT
# ─────────────────────────────────────────
@app.route('/pdf-split', methods=['POST'])
def pdf_split():
    if 'file' not in request.files:
        return jsonify({'error': 'ಫೈಲ್ ಕಂಡುಬಂದಿಲ್ಲ'}), 400
    file = request.files['file']
    pages_param = request.form.get('pages', '').strip()
    input_path = tmp_path('pdf')
    file.save(input_path)
    try:
        reader = PdfReader(input_path)
        total = len(reader.pages)
        pages_to_extract = parse_pages(pages_param, total)
        writer = PdfWriter()
        for i in pages_to_extract:
            writer.add_page(reader.pages[i])
        buf = io.BytesIO()
        writer.write(buf)
        buf.seek(0)
        return send_file(buf, mimetype='application/pdf', as_attachment=True, download_name='split.pdf')
    except Exception as e:
        return jsonify({'error': str(e)}), 500
    finally:
        cleanup(input_path)


# ─────────────────────────────────────────
# 5. IMAGE TO PDF
# ─────────────────────────────────────────
@app.route('/img-to-pdf', methods=['POST'])
def img_to_pdf():
    files = request.files.getlist('files')
    if not files:
        return jsonify({'error': 'ಚಿತ್ರಗಳು ಕಂಡುಬಂದಿಲ್ಲ'}), 400

    saved = []
    try:
        buf = io.BytesIO()
        c = canvas.Canvas(buf, pagesize=A4)
        w, h = A4

        for file in files:
            path = tmp_path('img')
            file.save(path)
            saved.append(path)

            img = Image.open(path)
            # Fix transparency issues
            if img.mode in ('RGBA', 'P', 'LA'):
                background = Image.new('RGB', img.size, (255, 255, 255))
                if img.mode == 'P':
                    img = img.convert('RGBA')
                background.paste(img, mask=img.split()[-1] if img.mode == 'RGBA' else None)
                img = background
                img.save(path, 'JPEG', quality=95)
            elif img.mode != 'RGB':
                img = img.convert('RGB')
                img.save(path, 'JPEG', quality=95)

            img_w, img_h = img.size
            ratio = min(w / img_w, h / img_h)
            new_w, new_h = img_w * ratio, img_h * ratio
            x = (w - new_w) / 2
            y = (h - new_h) / 2
            c.drawImage(path, x, y, new_w, new_h, preserveAspectRatio=True)
            c.showPage()

        c.save()
        buf.seek(0)
        return send_file(buf, mimetype='application/pdf', as_attachment=True, download_name='images.pdf')
    except Exception as e:
        return jsonify({'error': str(e)}), 500
    finally:
        cleanup(*saved)


# ─────────────────────────────────────────
# 6. PDF COMPRESS
# ─────────────────────────────────────────
@app.route('/pdf-compress', methods=['POST'])
def pdf_compress():
    if 'file' not in request.files:
        return jsonify({'error': 'ಫೈಲ್ ಕಂಡುಬಂದಿಲ್ಲ'}), 400
    file = request.files['file']
    if not file.filename.lower().endswith('.pdf'):
        return jsonify({'error': 'PDF ಮಾತ್ರ ಸ್ವೀಕಾರ'}), 400

    input_path = tmp_path('pdf')
    file.save(input_path)
    original_size = os.path.getsize(input_path)

    try:
        reader = PdfReader(input_path)
        writer = PdfWriter()

        for page in reader.pages:
            # Compress page contents
            page.compress_content_streams()
            writer.add_page(page)

        # Copy metadata
        if reader.metadata:
            writer.add_metadata(reader.metadata)

        # Enable object streams for better compression
        writer._header = b'%PDF-1.4'

        buf = io.BytesIO()
        writer.write(buf)
        buf.seek(0)

        compressed_size = buf.getbuffer().nbytes
        saved_percent = round((1 - compressed_size / original_size) * 100, 1)

        response = send_file(
            buf,
            mimetype='application/pdf',
            as_attachment=True,
            download_name='compressed.pdf'
        )
        # Send size info in headers
        response.headers['X-Original-Size'] = str(original_size)
        response.headers['X-Compressed-Size'] = str(compressed_size)
        response.headers['X-Saved-Percent'] = str(saved_percent)
        response.headers['Access-Control-Expose-Headers'] = 'X-Original-Size, X-Compressed-Size, X-Saved-Percent'
        return response

    except Exception as e:
        return jsonify({'error': str(e)}), 500
    finally:
        cleanup(input_path)


# ─────────────────────────────────────────
# HELPER
# ─────────────────────────────────────────
def parse_pages(pages_param, total):
    if not pages_param:
        return list(range(total))
    pages = set()
    for part in pages_param.split(','):
        part = part.strip()
        if '-' in part:
            s, e = part.split('-')
            for p in range(int(s) - 1, int(e)):
                if 0 <= p < total: pages.add(p)
        else:
            p = int(part) - 1
            if 0 <= p < total: pages.add(p)
    return sorted(pages)


if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    app.run(host='0.0.0.0', port=port)